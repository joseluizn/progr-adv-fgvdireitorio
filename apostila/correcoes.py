import sys
import os
import re
import requests

def aula2_ex1(name_space):
	#é importante passar o name space, porque assim eu não preciso passar cada uma das variáveis que serão
	#testadas em cada um dos exercícios
	#testa para ver se o aluno definiu as variáveis corretamente
	for i in ["contratante", "contratado", "honorarios_mes", "servicos_prestados", "honorarios_dia"]:
		if i not in name_space:
			print(f"Você se lembrou de definir a variável {i}?\nSerá que você cometeu algum erro de digitação?")
			return

	#recria internamente as variáveis para aumentar a legibilidade do código
	contratante = name_space["contratante"]
	contratado = name_space["contratado"]
	honorarios_mes = name_space["honorarios_mes"]
	honorarios_dia = name_space["honorarios_dia"]
	servicos_prestados = name_space["servicos_prestados"]

	#agora é preciso checar pelos tipos de cada variável
	for i in [contratante, contratado]:
		if type(i) != str:
			print("As variáveis 'contratante' e 'contratado' devem conter nomes. Lembre-se disso. Qual o tipo de variável adequado?\nCaso necessário, reveja as porções relevantes da aula acima")
			return

	if isinstance(honorarios_mes, (int, float)) == False:
		print("A variável honorarios_mes precisa ser um número.")
		return

	if honorarios_dia != (honorarios_mes / 30):
		print("A variável honorarios_dia precisa ser um número. Dica: ela precisa ser o resultado de uma divisão")
		return

	if type(servicos_prestados) != bool:
		print("A variável servicos_prestados precisa ser booleana.")
		return

	print("Parabéns! Você aprendeu sobre variáveis e tipos de dados em Python!")
	return

def aula3_ex1(item):
	if type(item) != int:
		print("Quando acessamos uma variável pelo índice, precisamos usar um inteiro!")
		return
	elif item == 6:
		print("Lembre-se, o Python começa a contar com 0! Começando a contar do 0, qual é o sexto elemento?")
		return
	elif item != 6 and item != 5:
		print("Queremos acessar o sexto item da lista")
		return
	elif item == 5:
		print("Parabéns! Sinal de que está prestando atenção!")

def aula3_ex2(name_space):
	if "condutas_petrechos" not in name_space:
		print("Você se lembrou de criar uma lista chamada 'condutas_petrechos'? É importante que o nome da lista seja exatamente esse!")
		return

	lista_condutas = name_space["condutas_petrechos"]

	if type(lista_condutas) != list:
		print("'condutas_petrechos' deve ser uma lista! Reveja seu código e, se necessário, a parte relevante da apostila. Listas são criadas entre [] e devem ter seus itens separados por ','.")
		return

	for i in ["fabricar", "adquirir", "fornecer", "possuir", "guardar"]:
		if i not in lista_condutas:
			print(f"Será que você se esqueceu de adicionar a conduta '{i}' à sua lista?")
			return

	if len(lista_condutas) != 5:
		print(f"O art. 294 especifica 5 condutas, mas sua lista contém {len(lista_condutas)} items...")
		return

	if "conduta_agente" not in name_space:
		print("O programa precisa avaliar a conduta do agente, mas não encontramos 'conduta_agente' dentre as suas variávels. Reveja seu programa! Ele só vai retornar o resultado correto se estiver avaliando uma variável chamada 'conduta_agente'.")
		return

	conduta = name_space["conduta_agente"]

	if type(conduta) != str:
		print("A conduta deve ser uma string, afinal, precisamos checar se ela consta de uma lista de strings.")
		return

	with open('exercicio.txt', 'r') as f:
		out_txt = f.read()

	if len(out_txt) == 0:
		print("Você se lembrou de pedir para que o programa imprimi-se (com print) o resultado?")
		return

	print(out_txt)

	if conduta in lista_condutas:
		if "O agente está sujeito à pena de reclusão de um a três anos e multa" not in out_txt:
			print("Essa não era a resposta que estávamos esperando... Lembre-se, o texto impresso tem que ser exatamente aquele colocado no enunciado.")
			return
		elif "O agente está sujeito à pena de reclusão de um a três anos e multa" in out_txt:
			print("Parabéns! Você acertou o exercício!")
			return
	else:
		if "A conduta descrita não está tipificada no art. 294" not in out_txt:
			print("Essa não era a resposta que estávamos esperando... Lembre-se, o texto impresso tem que ser exatamente aquele colocado no enunciado.")
			return
		elif "A conduta descrita não está tipificada no art. 294" in out_txt:
			print("Parabéns! Você acertou o exercício!")

def aula4_ex1(dicionario):
	for cliente in ['Maria Berenice Dias', 'Pontes de Miranda', 'José Afonso da Silva', 'Aliomar Baleeiro', 'Guilherme Nucci', 'Nelson Hungria', 'Sobral Pinto', 'Evandro Lins e Silva']:
		if cliente not in dicionario.keys():
			print(f"Você se esqueceu de adicionar {cliente} ao seu dicionário!")
			return

	valores_unicos_reais = []

	for item in list(dicionario.values()):
		if type(item) != list:
			valores_unicos_reais.append(item)
		else:
			for inf_item in item:
				valores_unicos_reais.append(inf_item)

	if len(valores_unicos_reais) != 20:
		print(f"Estranho... Começamos com 20 processos, mas seu dicionário contabiliza {len(valores_unicos_reais)} valores. Por favor, reveja seu código!")
		return

	for processo in ['2002.766393', '2011.487453', '2004.314066', '1998.516307', '2011.491429', '2003.714034', '2000.695102', '2012.628045', '2017.462447', '1996.518281', '1997.416985', '2002.321557','2002.737116', '2011.655314', '1990.414436', '2012.369756','1996.773047', '1995.344760', '1993.328622', '2016.681535']:
		if processo not in valores_unicos_reais:
			print(f"Você se esqueceu de adicionar o processo {processo} ao seu dicionário!")

	if len(valores_unicos_reais) != 20:
		print(f"Estranho... Começamos com 20 processos, mas seu dicionário contabiliza {len(valores_unicos_reais)} valores. Por favor, reveja seu código!")
		return
	else:
		print("Parabéns! Você acertou o exercício!")
		return

def aula4_ex2(dicionario):
	if dicionario["Guilherme Nucci"]["2012.628045"]["Órgão julgador"] == "60ª Câmara Cível":
		print("Parabéns! Você entendeu como funciona o método update()!")
		return
	else:
		print("Há algo de errado com o seu dicionário... Você se lembrou de atualizar o órgão julgador do proceso 2012.628045?\nÉ necessário que você tenha digitado o nome da câmara cível exatamente como citado no enunciado")


def aula5_ex1(dict_contratos):

	gabarito = {'Maria Berenice Dias' : 20000,
                 'Jorge Miranda' : 2000,
                 'Pontes de Miranda' : 40000}

	for key, value in gabarito.items():
		gabarito[key] = value + (value * 0.0375)

	for key in dict_contratos:
		if dict_contratos[key] != gabarito[key]:
			print(f"Há algo de errado com o seu código... O valor corrigido do contrato de {key} é {gabarito[key]}, e não {dict_contratos[key]}")
			return
	
	print("Parabéns! Você acertou o exercício!")


def aula6_ex1(fun_clausula1, texto_contrato):
	texto_final = fun_clausula1(texto_contrato, "la", "la", "la", "la", "la", "la")
	if texto_final is None:
		print("Você se lembrou de retornar o resultado da sua função? Use 'return'.")
		return
	for tag in ["<'nome'>", "<'nacionalidade'>", "<'estado_civil'>", "<'cpf'>", "<'endereco'>", "<'valor_contrato'>"]:
		if tag in texto_final:
			print(f"Parece que você se esqueceu de substituir a tag {tag} no texto final... Por favor, corrija seu código.")
			return
		else:
			print("Parabéns! Você completou seu primeiro programa de document assembly!")
			return

def aula6_ex2(usuarios_dict):
	corrected_files = {}
	for file in os.listdir():
		if file[-15:] == "_declaracao.txt":
			with open(file, "r") as f:
				file_text = f.read()
			corrected_files[file[:-15]] = file_text

	if not corrected_files:
		print("Parece que você não criou nenhum arquivo terminando com '_declaracao.txt' na sua pasta local...")
		return

	if len(corrected_files) != len(usuarios_dict):
		print(f"Há algo de errado... O dicionário usuarios tem {len(usuarios_dict)} itens,\nmas você só criou {len(corrected_files)} arquivos de acordo com nossa convenção")
		return

	for file_name, text in corrected_files.items():
		for usuario, usuario_dict in usuarios_dict.items():
			if usuario in file_name:
				for atributo, valor in usuario_dict.items():
					if str(valor) not in text:
						print(f"Vocẽ se esqueceu de incluir a seguinte informação a respeito de {usuario}: {valor}")
						return

	print("Parabéns! Você conhece os fundamentos de document assembly!")
	return


def aula8_ex1(lambda_fun):
	if 'lambda' not in str(lambda_fun):
		print("Parece que sua função não foi construída usando uma expressão lambda...")
		return

	texto = "<'nome'> é um grande cara"

	texto_tratado = lambda_fun(texto, "stringtest")

	if "<'nome'>" in texto_tratado:
		print("A função não foi bem sucedida em substituir todas as tags do texto de teste")
		return

	if "stringtest" not in texto_tratado:
		print("A função não substituiu todas as tags pelo nome que usamos no teste")
		return

	print("Parabéns! Você já sabe como usar funções lambda!")
	return

def aula9_ex1(penas_base, penas_finais):
	penas_bench = map(lambda x: x + (x * 1/5), penas_base)
	penas_compare = list(zip(penas_bench, penas_finais))
	init_index = 1
	for pair in penas_compare:
		if pair[0] != pair[1]:
			print(f"O item número {init_index} da sua lista parece estar errado...\nEle está com o número {pair[1]}, quando deveria conter {pair[0]}")
			return
		init_index += 1

	print("Parabéns! Você consegue combinar map() com lambda!")

def aula9_ex2(gabarito, precos, vendas):
	if len(gabarito) != 5:
		print(f"Sua resposta tem informações a respeito de {len(gabarito)} meses, quando deveria ter 5...")
		return

	precos_e_vendas = list(zip(precos, vendas))
	receita_mes = []
	for item in precos_e_vendas:
	    receita_mes.append(item[0] * item[1])

	for item in gabarito:
		if len(item) != 3:
			print(f"O enunciado informa que preciamos de 3 informações ordenadas em uma tupla para cada mês.\nSua resposta, porém, contém {len(item)} itens em cada tupla.")
			return
		if item[0] not in precos:
			print(f"O primeiro item de cada tupla deveria ser um dos preços contidos em lista_precos.\n{item[0]}, porém, não é um deles.")
			return
		if item[1] not in receita_mes:
			print(f"O segundo item de cada tupla deveria ser a receita mensal bruta da empresa.\n{item[1]}, porém, não é um resultado válido.")
			return		
		if item[2] not in list(map(lambda x: x - (x * 0.2), receita_mes)):
			print(f"O terceiro item de cada tupla deveria ser a receita após dedução de impostos.\n{item[2]}, porém, não é um resultado válido.")
			return

	print("Parabéns! Você aprendeu bem o conteúdo desta aula!")


def aula10_ex1(texto, cit_dict):
	indices_processos_citados = []

	for proc in range(texto.count(".8.19")):
	    if not indices_processos_citados:
	        new_index = texto.find(".8.19")
	    else:
	        new_index = texto.find(".8.19", (max(indices_processos_citados) + 20), len(texto))
	    indices_processos_citados.append(new_index - 15)
	    
	citacoes_gab = {}
	    
	for indice in indices_processos_citados:
	    num_cnj = texto[indice : indice + 25]
	    infos = num_cnj.split(".")
	    citacoes_gab[num_cnj] = {"id" : infos[0],
	                             "ano" : infos[1],
	                             "cod_just" : infos[2],
	                             "cod_trib" : infos[3],
	                             "org_julg" : infos[4]}

	#de fato inicia a correção
	if cit_dict == citacoes_gab:
		print("Parabéns! Esse desafio é bem difícil e você se saiu muito bem!")
		return

	for proc_gab, info_dict_gab in citacoes_gab.items():
		if proc_gab not in cit_dict.keys():
			print(f"Será que você se esqueceu do processo {proc_gab}? Ele não consta das chaves de seu dicionário.")
			return
		for info_name, info_value in info_dict_gab.items():
			if info_name not in cit_dict[proc_gab].keys():
				print(f"O campo {info_name} deveria constar do seu dicionário sobre o processo {proc_gab}")
				return
			if info_value != cit_dict[proc_gab][info_name]:
				print(f"Há algo de errado com o campo {info_name} do processo {proc_gab}. O valor deveria ser {info_value}, mas consta como {cit_dict[proc_gab][info_name]}")
				return

	print("Parabéns! Esse desafio é bem difícil e você se saiu muito bem!")
	return

def aula11_ex1(reg_exp):
	google_url = "https://drive.google.com/uc?export=download&id="
	file_id = "1JJ4Qf1gt1o2tTryhUFuCfloqFBJYkqdC"

	tjrj_1_text = requests.get(google_url + file_id).content.decode(encoding = "latin9")

	results_aluno = reg_exp.findall(tjrj_1_text)

	results_gabarito = re.compile(r"\d{7}\W\d{2}\W\d{4}\W\d\W\d{2}\W\d{4}").findall(tjrj_1_text)

	if results_aluno == results_gabarito:
		print("Parabéns! Você já consegue extrair números CNJ de textos.")
	else:
		for r in results_aluno:
			if r not in results_gabarito:
				print(f"Não esperávamos que {r} estivesse entre os resultados...")
		for rg in results_gabarito:
			if rg not in results_aluno:
				print(f"Esperávamos que {rg} estivesse entre os resultados.")

def aula13_ex1(file_name):
	import docx
	document = docx.Document(file_name)

	all_paragraphs = []

	for p in document.paragraphs:
		all_paragraphs.append(p.text)

	whole_text = "".join(all_paragraphs)

	if "99" not in whole_text:
		print("O número de vara que indicamos (99) não consta do documento final")
		return
	elif "2019.30999203-04" not in whole_text:
		print("O número de processo que indicamos (2019.30989203-04) não consta do documento final")
		return
	elif "Tício Romano" not in whole_text:
		print("O nome do seu cliente, Tício Romano, está faltando na petição")
		return
	elif "28" not in whole_text:
		print("A data indicada (dia 28) está errada.")
		return
	elif "junho" not in whole_text.lower():
		print("Informamos que o mês da petição deveria ser junho, mas 'junho' não consta do seu documento")
		return
	elif "2020" not in whole_text:
		print("O ano de 2020 não consta do seu documento, mas deveria estar na data")
		return
	elif "Caio Mário" not in whole_text:
		print("O signatário da petição será o advogado Caio Mário, porém, seu nome não consta do documento")
		return
	elif "00001" not in whole_text:
		print("O advogado signatário possui a inscrição 00001 na OAB/RJ, mas esse número não consta do documento")
		return
	else:
		print("Parabéns! Você completou o exercício e já sabe como montar petições automaticamente usando Python!")

	####################questionarios###############

def questionario1():
	perguntas = {"Qual procedimento define as filas?\na) LIFO (Last in, first out)\nb) FIFO (First in, first out)\nc) Filas são estruturas de dados como listas e dicionários e podem incorporar qualquer procedimento" : "b", "João ingressa na fila de um dos caixas do supermercado.\nHá uma pessoa sendo atendida e mais ninguém à sua frente. Alguns\nminutos depois, chegam, nessa ordem, Maria e Joana.\nDado que estamos tratando de uma fila, em qual ordem eles serão atendidos?\na)João, Maria e Joana\nb)Joana, Maria e João\nc) Maria, João e Joana" : "a", "Como eu consigo acessar o primeiro item de uma lista?\na)Através do método .head()\nb)Através do índice 1\nc) Através do índice 0" : "c"}
	for pergunta, gabarito in perguntas.items():
		resposta = str(input(pergunta + "\n")).lower()
		while resposta not in ["a", "b", "c"]:
			print("a, b e c são as únicas respostas possíveis. Por favor, digite uma dessas letras e aperte enter"),
			resposta = str(input(pergunta + "\n")).lower()
		if resposta != gabarito:
			print(f"Infelizmente, {resposta} não era a resposta correta... A opção correta era a letra {gabarito}.")
		if resposta == gabarito:
			print("Parabéns! Você acertou!")

def questionario2():
	perguntas = {"Qual procedimento define as pilhas?\na) LIFO (Last in, first out)\nb) FIFO (First in, first out)\nc) Pilhas são estruturas de dados especialmente bem desenhadas para o direito, como pode ser visto a partir de sua adoção em cartórios no Brasil inteiro." : "a", "João ingressa na fila de um dos caixas do supermercado.\nHá uma pessoa sendo atendida e mais ninguém à sua frente. Alguns\nminutos depois, chegam, nessa ordem, Maria e Joana.\nDado que estamos tratando de uma pilha, em qual ordem eles serão atendidos?\na)João, Maria e Joana\nb)Joana, Maria e João\nc) Maria, João e Joana" : "b"}
	for pergunta, gabarito in perguntas.items():
		resposta = str(input(pergunta + "\n")).lower()
		while resposta not in ["a", "b", "c"]:
			print("a, b e c são as únicas respostas possíveis. Por favor, digite uma dessas letras e aperte enter"),
			resposta = str(input(pergunta + "\n")).lower()
		if resposta != gabarito:
			print(f"Infelizmente, {resposta} não era a resposta correta... A opção correta era a letra {gabarito}.")
		if resposta == gabarito:
			print("Parabéns! Você acertou!")

def questionario3():
	perguntas = {"A tabela1 tem 34.423 linhas, a tabela2 tem 12.343 linhas. Quantas linhas terá o resultado de um left join entre a tabela1 e a tabela2?\na) 12.343 linhas\nb) 34.423 linhas\nc) 22.080 linhas" : "b", "A tabela1 tem 34.423 linhas, a tabela2 tem 12.343 linhas. Quantas linhas terá o resultado de um inner join entre a tabela1 e a tabela2?\na) 12.343 linhas\nb) 34.423 linhas\nc) 22.080 linhas" : "a", "A tabela1 tem 34.423 linhas, a tabela2 tem 12.343 linhas. Quantas linhas terá o resultado de um right join entre a tabela1 e a tabela2?\na) 12.343 linhas\nb) 34.423 linhas\nc) 22.080 linhas" : "a"}
	for pergunta, gabarito in perguntas.items():
		resposta = str(input(pergunta + "\n")).lower()
		while resposta not in ["a", "b", "c"]:
			print("a, b e c são as únicas respostas possíveis. Por favor, digite uma dessas letras e aperte enter"),
			resposta = str(input(pergunta + "\n")).lower()
		if resposta != gabarito:
			print(f"Infelizmente, {resposta} não era a resposta correta... A opção correta era a letra {gabarito}.")
		if resposta == gabarito:
			print("Parabéns! Você acertou!")

def questionario4():
	opcoes = ["a) 0.337\n", "b) 0.663\n", "c) 40%\n"]
	gabarito = "b"
	for opcao in opcoes:
		print(opcao)
	resposta = str(input("Qual das opções acima é a correta? ")).lower()
	while resposta not in ["a", "b", "c"]:
		print("a, b e c são as únicas respostas possíveis. Por favor, digite uma dessas letras e aperte enter"),
		resposta = str(input("Qual das opções acima é a correta? ")).lower()
	if resposta != gabarito:
		print(f"Infelizmente, {resposta} não era a resposta correta... A opção correta era a letra {gabarito}.")
	if resposta == gabarito:
		print("Parabéns! Você acertou!")


################exercicios extra######################

def exercicio_extra1():
	perguntas = {"É possível somar duas strings?\na) Não, soma é uma operação reservada para tipos numéricos\nb) Sim! O resultado será uma string maior, contendo as strings somadas\nc) Sim, mas o resultado será um inteiro contendo o tamanho das duas strings" : "b", "Todas as operações que podem ser feitas com strings podem ser feitas com tipos numéricos?\na) Não, existem operações que são exclusivas de strings\nb) Sim. Tipos numéricos são muito mais flexíveis que strings e podemos fazer qualquer coisa com eles \n" : "a", "É possível avaliar expressões envolvendo dois tipos de dados diferentes?\na) Não. O Python só aceita operações com dados do mesmo tipo\nb) Sim, podemos fazer operações entre strings e booleanos\nc) Sim. Podemos fazer operações com ints e floats" : "c"}
	for pergunta, gabarito in perguntas.items():
		resposta = str(input(pergunta + "\n")).lower()
		while resposta not in ["a", "b", "c"]:
			print("a, b e c são as únicas respostas possíveis. Por favor, digite uma dessas letras e aperte enter"),
			resposta = str(input(pergunta + "\n")).lower()
		if resposta != gabarito:
			print(f"Infelizmente, {resposta} não era a resposta correta... A opção correta era a letra {gabarito}.")
		if resposta == gabarito:
			print("Parabéns! Você acertou!")

def exercicio_extra2_1(lista_clientes):
	gabarito = ['Bianca', 'Igor', 'Juliana', 'Ana Luiza', 'Marcelo', 'Larissa','Julia']
	if lista_clientes == gabarito:
		print("Parabéns! Você acertou o exercício")
	else:
		for i in gabarito:
			if i not in lista_clientes:
				print(f"{i} não está na sua lista, mas era um dos clientes potenciais.")
				return

	for i in lista_clientes:
		if i not in gabarito:
			print(f"{i} não era um cliente potencial. Porque {i} consta da sua lista?")
			return

def exercicio_extra2_2(lista_clientes):
	gabarito = ['Igor', 'Juliana', 'Ana Luiza', 'Larissa', 'Julia']
	if lista_clientes == gabarito:
		print("Parabéns! Você acertou o exercício")
	else:
		if "Marcelo" in lista_clientes:
			print("Você se esqueceu de remover Giulia.")
			return
		elif "Bianca" in lista_clientes:
			print("Você se esqueceu de remover Bianca.")
		return


def exercicio_extra3(dict_clientes):
	for cliente in dict_clientes:
		for attr in ["idade", "cpf", "valor_mensal"]:
			if attr not in dict_clientes[cliente].keys():
				print(f"Você se esqueceu de salvar o atributo {attr} do cliente {cliente}.")
				return

	print("Parabéns! Você acertou o exercício!")

def exercicio_extra4(dct_fin, lista_cobrancas):
	gabarito = []
	for cliente, infos in dct_fin.items():
		valor_total = infos["valor_total"]
		valor_pago = infos["valor_pago"]
		parcelas_faltantes = infos["parcelas_faltantes"]
		while valor_pago < valor_total:
			valor_parcela = (valor_total - valor_pago) / parcelas_faltantes
			parcelas_faltantes -= 1
			valor_pago += valor_parcela
			valor_faltante = valor_total - valor_pago
			gabarito.append(cliente + " deve pagar " + str(valor_parcela) + " ao escritório. Depois dessa parcela, ficarão faltando mais " + str(parcelas_faltantes) + " parcelas a pagar, totalizando " + str(valor_faltante) + " reais")

	if not lista_cobrancas:
		print("A lista 'cobrancas' precisa estar preenchida!")
		return

	respostas_v_gabarito = list(zip(lista_cobrancas, gabarito))
	for pair in respostas_v_gabarito:
		if pair[0] != pair[1]:
			print(f"Em uma das linhas, sua resposta foi:\n{pair[0]}.\nQuando o correto era:\n {pair[1]}\nLembre-se, todas as vírgulas e espaços precisam estar corretas!")
			return

	print("Parabéns! Você compreendeu bem os loops for e while!")

def exercicio_extra5(progressao_anual):
	presos_dict_gab = {"Preso 1" : {"pena_total" : 15,
                           "pena_regime_atual" : 1,
                           "pena_cumprida_geral" : 1,
                           "regime" : "fechado"},
              "Preso 2" : {"pena_total" : 21,
                           "pena_regime_atual" : 3,
                           "pena_cumprida_geral" : 3,
                          "regime" : "fechado"},
              "Preso 3" : {"pena_total" : 7,
                           "pena_regime_atual" : 0,
                           "pena_cumprida_geral" : 0,
                          "regime" : "semi-aberto"},
              "Preso 4" : {"pena_total" : 2,
                          "pena_regime_atual" : 0,
                          "pena_cumprida_geral" : 0,
                          "regime" : "aberto"}}

	presos_dict_test = {"Preso 1" : {"pena_total" : 15,
                           "pena_regime_atual" : 1,
                           "pena_cumprida_geral" : 1,
                           "regime" : "fechado"},
              "Preso 2" : {"pena_total" : 21,
                           "pena_regime_atual" : 3,
                           "pena_cumprida_geral" : 3,
                          "regime" : "fechado"},
              "Preso 3" : {"pena_total" : 7,
                           "pena_regime_atual" : 0,
                           "pena_cumprida_geral" : 0,
                          "regime" : "semi-aberto"},
              "Preso 4" : {"pena_total" : 2,
                          "pena_regime_atual" : 0,
                          "pena_cumprida_geral" : 0,
                          "regime" : "aberto"}}

	def gabarito(presos):
		for key, preso in presos.items():
			if preso["pena_cumprida_geral"] < preso["pena_total"]:
				if preso["pena_regime_atual"] / preso["pena_total"] <= 1/6:
					preso["pena_regime_atual"] += 1
					preso["pena_cumprida_geral"] += 1
				else:
					if preso["regime"] == "fechado":
						preso["regime"] = "semi-aberto"
						preso["pena_regime_atual"] = 0
					elif preso["regime"] == "semi-aberto":
						preso["regime"] = "aberto"
						preso["pena_regime_atual"] = 0
					else:
						preso["pena_regime_atual"] += 1
					preso["pena_cumprida_geral"] += 1
			else:
				preso["regime"] = "pena cumprida"

		return presos

	gabarito_results = gabarito(presos_dict_gab)
	gabarito_results = gabarito(gabarito_results)
	gabarito_results = gabarito(gabarito_results)

	test = progressao_anual(presos_dict_test)
	test = progressao_anual(test)
	test = progressao_anual(test)

	for key, attrs in test.items():
		for attr, value in attrs.items():
			if value != gabarito_results[key][attr]:
				print(f"Há algo de errado com a sua resposta. O valor de {attr} para o {key} deveria ser {gabarito_results[key][attr]} após dois anos, e não {value}")
				return

	print("Parabéns! Você acertou o exercício extra!")

def exercicio_extra6(onde_estou):
	test = onde_estou()
	if os.getcwd() not in test:
		print(f"Seu diretório atual {os.getcwd()} não consta do resultado da função onde_estou.")
		return
	elif "Atualmente, você está no seguinte diretório:" not in test:
		print("Sua mensagem padrão é diferente daquela que pedimos.")
		return
	else:
		print("Parabéns! Você acertou o desafio extra!")

def exercicio_extra7(lista):
	gabarito = list(map(lambda x: x * 4, [2, 4, 8, 16, 32, 64]))

	if len(lista) != len(gabarito):
		print(f"Sua lista deveria conter 6 elementos (2, 4, 8, 16, 32, 64). Ao invés disso, ela contém {len(lista)} itens.")
		return

	for i, v in enumerate(lista):
		if v != gabarito[i]:
			print(f"{gabarito[i] / 4} vezes 4 = {gabarito[i]}, não {v}.")
			return

	print("Parabéns! Você acertou o exercício extra.")

def exercicio_extra8(paragrafos, text):
	gabarito = {}
	gabarito_list = text.split("\n")
	i = 1
	for p in gabarito_list:
		gabarito[i] = (len(p), p.count("consum"))
		i += 1

	for k, v in paragrafos.items():
		if v != gabarito[k]:
			print(f"Há algo de errado em seu código. Esperávamos a tupla {gabarito[k]} associada ao {k}º parágrafo.")
			return

	print("Parabéns! Você acertou o exercício extra.")

def exercicio_extra9(regex_extra):
	decisao = """
	Decisão: O Supremo Tribunal Federal reconheceu a repercussão geral da controvérsia objeto dos presentes autos que versa sobre o direito, ou não, do segurado contribuinte da Previdência Social, à luz do art. 5, XXXVI, da Constituição Federal, ao cálculo da aposentadoria de acordo com a legislação vigente à época em que preenchidos os requisitos exigidos para a sua concessão, a qual se revela mais vantajosa do que aquela vigente à data da efetiva jubilação. Trata-se do Tema n. 334 da Gestão por Temas da Repercussão Geral do STF, que será submetido à apreciação do Pleno desta Corte, nos autos do RE n. 630.501/RS, Relatora a Ministra Ellen Gracie.

	In casu, o acórdão recorrido assentou:

	“APOSENTADORIA. DIREITO ADQUIRIDO. LEI N 7.787, DE 1989. INEXISTÊNCIA DE PREJUÍZO.

	Inexiste direito adquirido do segurado de ver recalculada a renda mensal inicial de sua aposentadoria segundo as regras anteriores à Lei n 7.787, de 1989, sob as quais teria completado os requisitos do benefício, uma vez que a nova lei não trouxe nenhum prejuízo aos segurados” (fls. 86).

	Destarte, aplicando a decisão Plenária no RE n. 579.431, secundada, a posteriori pelo AI n. 503.064-AgR-AgR, Relator o Ministro CELSO DE MELLO; AI n. 811.626-AgR-AgR, Relator o Ministro RICARDO LEWANDOWSKI, e RE n. 513.473-ED, Relator o Ministro CEZAR PELUSO, determino a devolução dos autos ao Tribunal de origem (artigo 328, parágrafo único, do RISTF c.c. artigo 543-B e seus parágrafos do Código de Processo Civil).

	Publique-se.
	"""

	results_aluno = reg_exp.findall(decisao)

	results_gabarito = re.compile(r"\bRE\s\w\W\s\d*\W\d*|\bAI\s\w\W\s\d*\W\d*").findall(decisao)

	if results_aluno == results_gabarito:
		print("Parabéns! Você já consegue extrair precedentes de decisões do STF.")
	else:
		for r in results_aluno:
			if r not in results_gabarito:
				print(f"Não esperávamos que {r} estivesse entre os resultados...")
		for rg in results_gabarito:
			if rg not in results_aluno:
				print(f"Esperávamos que {rg} estivesse entre os resultados.")

def exercicio_extra10():
	import docx
	petitions_list = [
    {"{N_VARA}" : "99",
     "{N_PROCESSO}" : "2019.30989203-04",
     "{NOME_CLIENTE}" : "Tício Romano",
     "{DIA_DE_HOJE}" : "28",
     "{MES_ATUAL}" : "junho",
     "{ANO_ATUAL}" : "2020",
     "{NOME_ADVOGADO}" : "Caio Mário",
     "{N_OAB}" : "00001"},
    {"{N_VARA}" : "33",
     "{N_PROCESSO}" : "2019.546897-04",
     "{NOME_CLIENTE}" : "Semprônio Siciliano",
     "{DIA_DE_HOJE}" : "29",
     "{MES_ATUAL}" : "junho",
     "{ANO_ATUAL}" : "2020",
     "{NOME_ADVOGADO}" : "Caio Mário Filho",
     "{N_OAB}" : "444444"},
    {"{N_VARA}" : "28",
     "{N_PROCESSO}" : "2019.8847214-01",
     "{NOME_CLIENTE}" : "Mévio Scipriano",
     "{DIA_DE_HOJE}" : "15",
     "{MES_ATUAL}" : "junho",
     "{ANO_ATUAL}" : "2020",
     "{NOME_ADVOGADO}" : "Caio Mário",
     "{N_OAB}" : "00001"},
    {"{N_VARA}" : "18",
     "{N_PROCESSO}" : "2019.111477-00",
     "{NOME_CLIENTE}" : "Semprônio Siciliano",
     "{DIA_DE_HOJE}" : "01",
     "{MES_ATUAL}" : "julho",
     "{ANO_ATUAL}" : "2020",
     "{NOME_ADVOGADO}" : "Caio Mário Filho",
     "{N_OAB}" : "444444"},
    {"{N_VARA}" : "99",
     "{N_PROCESSO}" : "2019.30989208-04",
     "{NOME_CLIENTE}" : "Justiniano Augusto",
     "{DIA_DE_HOJE}" : "28",
     "{MES_ATUAL}" : "junho",
     "{ANO_ATUAL}" : "2020",
     "{NOME_ADVOGADO}" : "Caio Mário",
     "{N_OAB}" : "00001"},
    {"{N_VARA}" : "43",
     "{N_PROCESSO}" : "2001.1007546-03",
     "{NOME_CLIENTE}" : "Tício Romano",
     "{DIA_DE_HOJE}" : "28",
     "{MES_ATUAL}" : "junho",
     "{ANO_ATUAL}" : "2020",
     "{NOME_ADVOGADO}" : "Caio Mário",
     "{N_OAB}" : "00001"}]
	
	for i, file in enumerate(["peticao01.docx", "peticao02.docx", "peticao03.docx", "peticao04.docx", "peticao05.docx", "peticao06.docx"]):
		document = docx.Document(file)

		all_paragraphs = []

		for p in document.paragraphs:
			all_paragraphs.append(p.text)
	
		whole_text = "".join(all_paragraphs)

		for tag, replacement in petitions_list[i].items():
			if replacement not in whole_text:
				print(f"Há algo de errado. Não encontramos a string {replacement} no documento {file}, que deveria ter substituído a tag {tag}.")
				return

	print("Parabéns! Você acertou o desafio extra!")